"""Document text extraction for PDF, DOCX, and TXT files."""

import os
import re

import PyPDF2
import docx


def _clean_text(text: str) -> str:
    """Normalize whitespace and strip junk characters from extracted text."""
    # Normalize unicode punctuation / non-breaking spaces
    text = (
        text.replace("\xa0", " ")
        .replace("\u2019", "'")
        .replace("\u2018", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2013", "-")
        .replace("\u2014", "-")
    )
    # Collapse runs of whitespace (but keep newlines)
    text = re.sub(r"[^\S\n]+", " ", text)
    # Collapse 3+ blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_file(file_path: str, filename: str) -> str:
    """
    Extract plain text from a file based on its extension.

    Args:
        file_path: Absolute path to the saved file on disk.
        filename:  Original filename (used to determine file type).

    Returns:
        Extracted text string, or an error message on failure.
    """
    try:
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            return _extract_pdf(file_path)
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return _clean_text(f.read())
        elif ext == ".docx":
            return _extract_docx(file_path)
        return f"Unsupported file type: '{ext}'"
    except Exception as e:
        return f"Error reading file: {e}"


def _extract_pdf(file_path: str) -> str:
    """Extract text from a PDF file page by page, labelling each page."""
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[Page {i + 1}]\n{text}")
            return _clean_text("\n\n".join(pages))
    except Exception as e:
        return f"Error reading PDF: {e}"


def _extract_docx(file_path: str) -> str:
    """Extract text from a DOCX file — paragraphs and tables."""
    try:
        doc = docx.Document(file_path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return _clean_text("\n".join(parts))
    except Exception as e:
        return f"Error reading DOCX: {e}"
